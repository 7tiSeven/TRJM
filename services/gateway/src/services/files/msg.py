"""
TRJM Gateway - MSG File Parser
===============================
Parser for Microsoft Outlook email files
"""

import io
from dataclasses import dataclass
from typing import List, Optional

import extract_msg

from ...core.logging import logger
from .parser import FileParser, Paragraph, ParsedDocument, ParserRegistry


@dataclass
class EmailContent:
    """Structured email content."""

    subject: str
    sender: str
    recipients: List[str]
    date: Optional[str]
    body: str
    has_attachments: bool
    attachment_names: List[str]


class MsgParser(FileParser):
    """
    Parser for .msg files (Outlook emails).

    Extracts email headers and body. Attachments are ignored
    for security reasons.
    """

    @property
    def supported_extension(self) -> str:
        return ".msg"

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse MSG file.

        Extracts subject, headers, and body.
        """
        # Parse MSG file
        msg = extract_msg.Message(io.BytesIO(content))

        try:
            # Extract email data
            subject = msg.subject or ""
            sender = msg.sender or ""
            recipients = msg.to or ""
            date = msg.date or ""
            body = msg.body or ""

            # Get attachment info (names only, not content)
            attachment_names = [att.longFilename or att.shortFilename for att in msg.attachments]
            has_attachments = len(attachment_names) > 0

            # Build paragraphs for translation
            paragraphs = []
            para_index = 0

            # Subject as first paragraph
            if subject:
                paragraphs.append(
                    Paragraph(
                        text=subject,
                        index=para_index,
                        metadata={"type": "subject"},
                    )
                )
                para_index += 1

            # Body paragraphs
            if body:
                body_paras = body.split("\n\n")
                for para_text in body_paras:
                    stripped = para_text.strip()
                    if stripped:
                        paragraphs.append(
                            Paragraph(
                                text=stripped,
                                index=para_index,
                                metadata={"type": "body"},
                            )
                        )
                        para_index += 1

            # Build full content for translation
            full_content = f"Subject: {subject}\n\n{body}"

            metadata = {
                "subject": subject,
                "sender": sender,
                "recipients": recipients,
                "date": date,
                "has_attachments": has_attachments,
                "attachment_names": attachment_names,
                "attachment_count": len(attachment_names),
            }

            logger.debug(
                "MSG parsed",
                filename=filename,
                subject=subject[:50] if subject else None,
                body_length=len(body),
                attachments=len(attachment_names),
            )

            return ParsedDocument(
                content=full_content,
                paragraphs=paragraphs,
                metadata=metadata,
                format_hints={
                    "is_email": True,
                    "has_attachments": has_attachments,
                },
                file_type="msg",
                file_name=filename,
                file_size=len(content),
            )

        finally:
            msg.close()

    async def generate(
        self,
        original: ParsedDocument,
        translated_paragraphs: List[Paragraph],
    ) -> bytes:
        """
        Generate translated output.

        Since we can't create MSG files easily, we output a DOCX summary
        with the translated email content.
        """
        from docx import Document

        doc = Document()

        # Add email header info
        doc.add_heading("Translated Email", 0)

        # Original metadata
        meta = original.metadata
        doc.add_paragraph(f"From: {meta.get('sender', 'Unknown')}")
        doc.add_paragraph(f"To: {meta.get('recipients', 'Unknown')}")
        doc.add_paragraph(f"Date: {meta.get('date', 'Unknown')}")

        if meta.get("has_attachments"):
            attachments = meta.get("attachment_names", [])
            doc.add_paragraph(f"Attachments: {', '.join(attachments)} (not translated)")

        doc.add_paragraph()

        # Translated subject
        subject_paras = [p for p in translated_paragraphs if p.metadata.get("type") == "subject"]
        if subject_paras:
            doc.add_heading("Subject", level=1)
            doc.add_paragraph(subject_paras[0].text)

        # Translated body
        body_paras = [p for p in translated_paragraphs if p.metadata.get("type") == "body"]
        if body_paras:
            doc.add_heading("Body", level=1)
            for para in body_paras:
                doc.add_paragraph(para.text)

        # Footer
        doc.add_paragraph()
        doc.add_paragraph("---")
        footer = doc.add_paragraph()
        footer.add_run("Translated by TRJM Agentic AI Translator").italic = True

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output.read()

    def get_output_filename(self, original_filename: str, target_lang: str) -> str:
        """MSG files are output as DOCX."""
        from pathlib import Path

        path = Path(original_filename)
        return f"{path.stem}_{target_lang}.docx"


# Register parser
ParserRegistry.register(MsgParser())
