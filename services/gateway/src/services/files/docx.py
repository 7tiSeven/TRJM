"""
TRJM Gateway - DOCX File Parser
================================
Parser for Microsoft Word documents
"""

import io
from typing import List

from docx import Document
from docx.shared import Inches, Pt

from ...core.logging import logger
from .parser import FileParser, Paragraph, ParsedDocument, ParserRegistry


class DocxParser(FileParser):
    """Parser for .docx files."""

    @property
    def supported_extension(self) -> str:
        return ".docx"

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX file.

        Extracts paragraphs while preserving structure hints.
        """
        doc = Document(io.BytesIO(content))

        paragraphs = []
        full_text_parts = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # Capture paragraph style for later reconstruction
                style_name = para.style.name if para.style else "Normal"

                # Check for formatting
                has_bold = any(run.bold for run in para.runs)
                has_italic = any(run.italic for run in para.runs)

                paragraphs.append(
                    Paragraph(
                        text=text,
                        index=i,
                        metadata={
                            "style": style_name,
                            "has_bold": has_bold,
                            "has_italic": has_italic,
                            "alignment": str(para.alignment) if para.alignment else None,
                        },
                    )
                )
                full_text_parts.append(text)

        # Extract document metadata
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title,
            "author": core_props.author,
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
        }

        # Note: Tables are not translated in this implementation
        # They could be added as a future enhancement

        logger.debug(
            "DOCX parsed",
            filename=filename,
            paragraphs=len(paragraphs),
            tables=len(doc.tables),
        )

        return ParsedDocument(
            content="\n\n".join(full_text_parts),
            paragraphs=paragraphs,
            metadata=metadata,
            format_hints={
                "has_tables": len(doc.tables) > 0,
                "has_images": False,  # Not checking for images currently
            },
            file_type="docx",
            file_name=filename,
            file_size=len(content),
        )

    async def generate(
        self,
        original: ParsedDocument,
        translated_paragraphs: List[Paragraph],
    ) -> bytes:
        """
        Generate translated DOCX file.

        Creates a new document preserving basic structure.
        """
        doc = Document()

        # Set document properties
        doc.core_properties.title = f"Translated: {original.metadata.get('title', '')}"

        # Add translated paragraphs
        for para in translated_paragraphs:
            p = doc.add_paragraph()

            # Apply style if available
            style = para.metadata.get("style", "Normal")
            try:
                p.style = style
            except KeyError:
                p.style = "Normal"

            # Add text with basic formatting
            run = p.add_run(para.text)

            # Apply formatting from original
            if para.metadata.get("has_bold"):
                run.bold = True
            if para.metadata.get("has_italic"):
                run.italic = True

        # Add note about translation
        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run("---").italic = True
        note = doc.add_paragraph()
        note.add_run("Translated by TRJM Agentic AI Translator").italic = True

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output.read()


# Register parser
ParserRegistry.register(DocxParser())
